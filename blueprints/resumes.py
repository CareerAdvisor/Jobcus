import base64, re, json, logging, os
from io import BytesIO

from flask import Blueprint, request, jsonify, current_app, make_response, send_file, render_template, Response
from flask_login import login_required, current_user

from limits import feature_enabled, check_and_increment
from abuse_guard import allow_free_use  # NEW: device/user-scoped guard
from auth_utils import api_login_required

import docx
from weasyprint import HTML, CSS
from docxtpl import DocxTemplate
from PyPDF2 import PdfReader
from jinja2 import TemplateNotFound
from authz import require_plan
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# OpenAI SDKs have changed slightly across versions; this keeps RateLimitError optional.
try:
    from openai import RateLimitError  # v0.x and v1.x expose this
except Exception:  # pragma: no cover
    class RateLimitError(Exception):  # fallback so except still works
        pass

resumes_bp = Blueprint("resumes", __name__)

# -------------------------------------------------------------------
# Print/PDF overrides (reduce WeasyPrint warnings, enforce margins)
# -------------------------------------------------------------------
PDF_CSS_OVERRIDES = """
@page { size: A4; margin: 0.75in; }    /* 0.75" on all sides */
* { box-shadow: none !important; }     /* WeasyPrint ignores box-shadow; silence visually */
@media print {
  html, body { background: white !important; }
  .resume-page, .resume { box-shadow: none !important; }
  .section { margin-top: 12px !important; }
  .section .rule { margin: 6px 0 10px 0 !important; }
  h1, h2, h3 { page-break-after: avoid; }
  .item-header { page-break-inside: avoid; }
  ul.bullets { margin-top: 6px !important; }
}
"""

# ---------- Helper: fallback context if OpenAI is unavailable ----------
def naive_context(data: dict) -> dict:
    import re as _re
    first = (data.get("firstName") or "").strip()
    last  = (data.get("lastName") or "").strip()
    derived = " ".join([p for p in (first, last) if p]).strip()

    name      = (data.get("name") or data.get("fullName") or derived or "").strip()
    title     = (data.get("title") or "").strip()
    contact   = (data.get("contact") or "").strip()
    summary   = (data.get("summary") or "").strip()
    edu_txt   = (data.get("education") or "").strip()
    exp_txt   = (data.get("experience") or "").strip()
    skills_s  = (data.get("skills") or "")
    certs_s   = (data.get("certifications") or "")
    portfolio = (data.get("portfolio") or "").strip()

    skills = [s.strip() for s in str(skills_s).replace("\n", ",").split(",") if s.strip()]
    certifications = [c.strip() for c in str(certs_s).splitlines() if c.strip()]

    education = []
    if edu_txt:
        education = [{
            "degree": edu_txt, "school": "", "location": "", "graduated": ""
        }]

    experience = []
    if exp_txt:
        bullets = [_re.sub(r'^[\-\•]\s*', '', ln).strip()
                   for ln in exp_txt.splitlines() if ln.strip()]
        experience = [{
            "role": title or "Experience",
            "company": "", "location": "",
            "start": "", "end": "",
            "bullets": bullets
        }]

    links = [{"url": portfolio, "label": "Portfolio"}] if portfolio else []

    return {
        "name": name, "title": title, "contact": contact, "summary": summary,
        "skills": skills, "links": links,
        "experience": experience, "education": education,
        "certifications": certifications,
    }

# ---------- Helper: normalize context coming from the UI ----------
def _normalize_ctx(data: dict) -> dict:
    import re as _re
    first = (data.get("firstName") or "").strip()
    last  = (data.get("lastName") or "").strip()
    derived = " ".join([p for p in (first, last) if p]).strip()
    name = (data.get("name") or data.get("fullName") or data.get("full_name") or derived or "").strip()

    ctx = dict(data)  # shallow copy
    ctx["name"] = name
    ctx["fullName"] = name
    ctx["full_name"] = name

    # Skills: allow CSV/newlines or list
    skills = ctx.get("skills") or []
    if isinstance(skills, str):
        skills = [s.strip() for s in _re.split(r"[,\n]", skills) if s.strip()]
    ctx["skills"] = skills

    # Certifications: allow newlines or list
    certs = ctx.get("certifications") or []
    if isinstance(certs, str):
        certs = [c.strip() for c in certs.splitlines() if c.strip()]
    ctx["certifications"] = certs

    # Ensure arrays exist
    ctx["experience"] = ctx.get("experience") or []
    ctx["education"]  = ctx.get("education")  or []
    ctx["links"]      = ctx.get("links")      or []

    return ctx

# ---------- Smart picker: normalize first, fall back to naive ----------
def build_context(data: dict) -> dict:
    data = data or {}
    ctx = _normalize_ctx(data)
    if not ctx.get("experience") and not ctx.get("education"):
        return naive_context(data)
    return ctx


# === ATS model: weights and deterministic checks ===
import math, datetime, collections, statistics, itertools

# We’ll use this to approximate pages for text resumes (PDF page count is not reliable post-OCR)
WORDS_PER_PAGE = 600

MONTHS = r"(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t|tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
DATE_PAT = re.compile(
    rf"\b{MONTHS}\s+\d{{4}}\b|\b\d{{1,2}}/\d{{4}}\b|\b(20\d{{2}}|19\d{{2}})\b",
    re.I
)  # "January 2021", "01/2021", or "2021"

EMAIL_PAT = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)
PHONE_PAT = re.compile(r"\+?\d[\d\-\s()]{7,}\d")
LOC_PAT   = re.compile(r"\b([A-Z][a-z]+(?:[ ,][A-Z][a-z]+)*)\b")  # very light

# Weak/buzzwords and action verbs
ACTION_VERBS_STRONG = {
    "achieved","analyzed","built","configured","created","delivered","designed","developed","drove",
    "implemented","improved","increased","launched","led","managed","optimized","owned","resolved",
    "reduced","shipped","streamlined","spearheaded","supported","automated","coordinated","deployed",
    "orchestrated","migrated","modernized","scaled","mentored"
}
WEAK_PHRASES = {
    "responsible for","helped","assisted","worked on","participated in","involved in"
}
BUZZWORDS = {
    "synergy","go-getter","outside the box","hard worker","team player","self-starter",
    "results-driven","rockstar","guru","ninja","dynamic","best-in-class","fast-paced"
}
PERSONAL_DATA = {"marital status","religion","nationality","date of birth","dob","photo","photograph"}

CERT_HINTS = {"aws","amazon web services","itil","pmp","cissp","cisa","az-","ms-","scrum","csm","ccna"}

BULLET_MARKERS = ("•","-","–","—","*")

# --- Keyword filtering & hints ----------------------------------------------
STOPWORDS_KW = {
    "a","an","and","the","of","for","with","from","into","onto","within","across",
    "about","above","below","under","over","at","by","to","in","on","as","is","are",
    "was","were","be","been","being","it","its","this","that","those","these","your",
    "you","our","their","my","we","or","nor","not","but","if","then","so","than",
    "very","more","most","can","will","would","should","could","may","might","etc",
    "currently","presently","strong","solid","excellent","good","great","private",
    "public","sector","known","having","combines","background","bridge","secure"
}
NOISE_VERBS = {
    "led","lead","manage","managed","own","owned","coordinate","coordinated","work",
    "worked","help","helped","assist","assisted","support","supported","drive","drove"
}
NOISE_NOUNS = {
    "experience","summary","objective","profile","responsibilities","duties",
    "team","teams","environment","company","organization","client","customers",
    "role","position","project","projects"
}

SKILL_HINTS = {
    "agile","scrum","kanban","waterfall","prince2","pmp","itil","change","risk",
    "stakeholder","stakeholders","governance","scope","budget","timeline","schedule",
    "roadmap","dependency","dependencies","ra id","raid","risk register","issue log",
    "project charter","business case","u at","uat","test plan","qa","kpi","okrs","okr",
    "jira","confluence","microsoft project","ms project","azure devops","power bi",
    "excel","visio","sharepoint","service now","servicenow","sql","api","sd lc","sdlc"
}
TITLE_HINTS  = {"project manager","delivery manager","scrum master","program manager"}
DOMAIN_HINTS = {"healthcare","fintech","ecommerce","telecom","banking","government","cloud","security","cybersecurity"}

def _is_acronym(w: str) -> bool:
    return 2 <= len(w) <= 6 and w.isupper()

# Sanitize keyword lists before returning to UI
def _clean_kw_list(lst):
    out = []
    for t in lst or []:
        t = t.strip().lower()
        parts = re.findall(r"[a-z0-9+\-/#]+", t)
        while parts and parts[0] in STOPWORDS_KW: parts.pop(0)
        while parts and parts[-1] in STOPWORDS_KW: parts.pop()
        if not parts:
            continue
        cleaned = " ".join(parts)

        # drop super short non-acronyms
        if (len(cleaned) < 3 and not _is_acronym(cleaned.upper())):
            continue

        # --- NEW: drop likely fragments/truncated stems like "projec", "qualit"
        if re.match(r"^[a-z]{3,}$", cleaned):
            # ends with two consonants or ends with a consonant after a long run can be OK,
            # but short stems missing common endings are often bad. Filter a few patterns:
            if re.search(r"(?:projec|qualit|requir|deliv|suppor|managemen|depar|operat|technolog)$", cleaned):
                continue

        if cleaned in STOPWORDS_KW or cleaned in NOISE_NOUNS or cleaned in NOISE_VERBS:
            continue

        out.append(cleaned)

    return sorted(set(out))

def _build_jd_terms(jd: str):
    """
    Extract meaningful JD terms:
    - remove stopwords/boilerplate/generic verbs
    - keep skills/tools/methods/titles/certs/domains (+ acronyms)
    - include useful bigrams/trigrams (e.g., 'risk register', 'ms project')
    """
    if not jd:
        return {"skills": [], "titles": [], "certs": [], "all": [], "acronyms": {}}

    words = re.findall(r"[A-Za-z][A-Za-z+\-/#0-9]*", jd)
    toks  = [w.lower() for w in words]

    bigrams  = [" ".join([toks[i], toks[i+1]]) for i in range(len(toks)-1)]
    trigrams = [" ".join([toks[i], toks[i+1], toks[i+2]]) for i in range(len(toks)-2)]

    singles = [
        t for t in toks
        if (_is_acronym(words[toks.index(t)]) or len(t) >= 3)
        and t not in STOPWORDS_KW
        and t not in NOISE_VERBS
        and t not in NOISE_NOUNS
    ]

    def looks_useful_phrase(p: str) -> bool:
        if any(h in p for h in (SKILL_HINTS | TITLE_HINTS | DOMAIN_HINTS)):
            return True
        if any(c in p for c in CERT_HINTS):
            return True
        if any(t in p for t in TITLE_HINTS):
            return True
        return False

    phrases = [p for p in (bigrams + trigrams) if looks_useful_phrase(p)]

    acronyms = {
        "aws": "amazon web services",
        "ad": "active directory",
        "sql": "structured query language",
        "k8s": "kubernetes",
        "mdm": "mobile device management",
        "sso": "single sign on",
        "u at": "user acceptance testing",
        "sd lc": "software development life cycle"
    }

    expanded = set(singles) | set(phrases)
    for a, full in acronyms.items():
        if a in expanded:   expanded.add(full)
        if full in expanded: expanded.add(a)

    certs  = sorted({w for w in expanded if any(c in w for c in CERT_HINTS)})
    titles = sorted({w for w in expanded if any(t in w for t in TITLE_HINTS)})
    skills = sorted({
        w for w in expanded
        if w not in set(certs) | set(titles)
        and (any(h in w for h in (SKILL_HINTS | DOMAIN_HINTS))
             or _is_acronym(w.replace(" ", "").upper()))
    })

    all_terms = sorted(set(skills) | set(titles) | set(certs), key=lambda x: (-len(x), x))
    return {"skills": skills, "titles": titles, "certs": certs, "all": all_terms, "acronyms": acronyms}

# --------------------------- helpers ---------------------------
def _clean_lines(text: str):
    return [ln.strip() for ln in text.splitlines() if ln.strip()]

def _tokenize(text: str):
    return re.findall(r"[A-Za-z][A-Za-z+\-/#0-9]*", text.lower())

def _split_sections(text: str):
    """Split into canonical sections so we can measure distribution."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    sections = collections.OrderedDict()
    cur = "_start"; sections[cur] = []

    def canonical(h: str) -> str:
        h = re.sub(r"[^a-z]+", " ", h.lower()).strip()
        if any(w in h for w in ("experience","work history","employment history","relevant experience","professional experience")):
            return "experience"
        if "education" in h:        return "education"
        if "skill" in h:            return "skills"
        if "project" in h:          return "projects"
        if "cert" in h or "license" in h: return "certifications"
        if "summary" in h or "profile" in h: return "summary"
        if "objective" in h:        return "objective"
        return h

    for ln in lines:
        if len(ln) <= 48 and re.match(r"^[A-Z][A-Za-z0-9 &/]+$", ln):
            key = canonical(ln)
            if key != "_start":
                cur = key
                sections[cur] = []
                continue
        sections[cur].append(ln)
    return sections

# --- Normalize text extracted from PDF/DOCX so words aren't split ---
def _normalize_extracted_text(s: str) -> str:
    """
    Fix common extractor artifacts:
      - remove soft/zero-width hyphens & BOMs
      - join hyphenated line breaks: 'projec-\n t' -> 'project'
      - join accidental mid-word breaks without hyphen: 'projec\n t' -> 'project'
      - collapse excessive whitespace
    """
    if not s:
        return s

    # 1) strip invisible/soft hyphens & zero-width chars
    s = re.sub(r"[\u00AD\u200B-\u200F\uFEFF]", "", s)

    # 2) hyphenated line breaks (ASCII or Unicode hyphens)
    #    e.g., end of line "projec-\n" then "t" on next line → join
    s = re.sub(r"([A-Za-z])[\-\u2010-\u2015]\s*\n\s*([A-Za-z])", r"\1\2", s)

    # 3) mid-word newline without hyphen: "projec\n t" → "project"
    s = re.sub(r"(?<=[A-Za-z])\s*\n\s*(?=[A-Za-z])", "", s)

    # (keep paragraph breaks when there are multiple blank lines)
    # replace 3+ newlines with two; collapse runs of spaces
    s = re.sub(r"\n{3,}", "\n\n", s)
    s = re.sub(r"[ \t]{2,}", " ", s)

    return s

def _extract_roles_with_dates(text: str):
    """Return list of roles with date hints + lines (to weight recency and detect gaps)."""
    sections = _split_sections(text)
    all_lines = list(itertools.chain.from_iterable(sections.values()))
    roles = []
    buf = []
    for ln in all_lines:
        if DATE_PAT.search(ln):
            if buf:
                roles.append({"lines": buf[:]})
                buf.clear()
        buf.append(ln)
    if buf:
        roles.append({"lines": buf[:]})

    now = datetime.date.today().year
    for r in roles:
        years = [int(y) for y in re.findall(r"\b(20\d{2}|19\d{2})\b", " ".join(r["lines"]))]
        r["start"] = (min(years) if years else None)
        r["end"]   = (now if re.search(r"\b(present|current)\b", " ".join(r["lines"]), re.I) else (max(years) if years else None))
    return roles

def _years_of_experience(roles):
    yrs = 0
    for r in roles:
        if r.get("start") and r.get("end"):
            yrs += max(0, r["end"] - r["start"])
    return yrs

def _has_action_verb_start(line: str):
    s = line
    s = re.sub(r"^(\*|•|-|\u2022|\u25CF)\s+", "", s)
    m = re.match(r"^([A-Za-z\-']+)", s)
    w = m.group(1).lower() if m else ""
    return w in ACTION_VERBS_STRONG, w

def _quant_stats(text: str):
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    bullets = [ln for ln in lines if ln[:2] in ("- ", "* ", "• ") or (ln and ln[0] in BULLET_MARKERS)]

    # Fallback: treat experience lines that start with a strong action verb as bullets
    if not bullets:
        ex = _split_sections(text).get("experience", [])
        pseudo = []
        for ln in ex:
            starts_with_verb, first = _has_action_verb_start(ln)
            if starts_with_verb and 6 <= len(ln.split()) <= 30:
                pseudo.append(ln)
        bullets = pseudo

    has_num = [
        bool(re.search(r"(\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b|\b\d+%|\b\d+\s?(?:k|m)\b|reduced|increased|saved|cut|grew|boosted)", b, re.I))
        for b in bullets
    ]
    starts  = [_has_action_verb_start(b)[0] for b in bullets]
    avg_len = statistics.mean([len(b.split()) for b in bullets]) if bullets else 0

    return {
        "bullet_count": len(bullets),
        "pct_with_numbers": int(round(100*sum(has_num)/max(1,len(bullets)))),
        "pct_action_starts": int(round(100*sum(starts)/max(1,len(bullets)))),
        "avg_words_per_bullet": round(avg_len,1)
    }

def _parse_and_format_score(resume_text: str, file_kind: str, raw_bytes: bytes|None):
    # 1) selectable text
    selectable = len(resume_text.strip()) >= 180
    s1 = 4 if selectable else 0

    # 2) file type
    s2 = 2 if file_kind in ("docx","pdf") else 0

    # 3) parsing traps: tables/columns/headers—heuristics
    txt = resume_text
    traps = 0
    if re.search(r"\|.+\|", txt): traps += 1               # ascii tables
    if txt.count("\t") > 30: traps += 1                    # tabular layout
    if re.search(r"\s{4,}\S", txt): traps += 1             # multi-column alignment
    s3 = max(0, 3 - traps)

    # 4) standard bullets/fonts (we can only check bullets)
    s4 = 1 if re.search(r"(^|\n)\s*(?:•|\-|\*)\s+", txt) else 0

    # image-only PDF heuristic (bytes >> text)
    hard_fail = False
    if raw_bytes and file_kind == "pdf":
        bpc = len(raw_bytes) / max(1, len(txt.encode("utf-8")))
        if selectable is False or bpc > 60:
            hard_fail = True

    return (0 if hard_fail else (s1+s2+s3+s4)), hard_fail

def _sections_structure_score(text: str):
    sec = _split_sections(text)
    has_std = {
        "experience": any("experience" in k for k in sec.keys()),
        "education": any("education" in k for k in sec.keys()),
        "skills": any("skills" in k for k in sec.keys()),
        "certs": any("certification" in k for k in sec.keys()),
    }
    base = 0
    base += 2 if EMAIL_PAT.search(text) else 0
    base += 2 if PHONE_PAT.search(text) else 0
    base += 1 if LOC_PAT.search(text) else 0

    roles = _extract_roles_with_dates(text)
    dated   = sum(1 for r in roles if r.get("start") or r.get("end"))
    chrono  = 1 if len(roles) < 2 else int(all((roles[i].get("end") or 0) >= (roles[i+1].get("end") or 0) for i in range(len(roles)-1)))
    rc = min(5, dated) + chrono  # max 6 for this part

    return {
        "score": (8 * sum(has_std.values())/4) + min(5, rc) + base,  # cap at 15
        "std": has_std,
        "reverse_chrono": bool(chrono),
        "roles_parsed": len(roles)
    }

# ---------------------- JD keyword scoring (0–100) ----------------------
def _keyword_match_to_jd(resume_text: str, jd_terms: dict, roles: list):
    """
    Bucketed keyword scoring → normalized to 0–100.
      Buckets (max points 35 total):
        - Skills/tools coverage .......... 15
        - Role/domain/title terms ........  8
        - Certifications/quals ...........  6
        - Acronym ↔ full form coverage ...  3
        - Natural distribution ...........  3
    """
    txt = resume_text.lower()
    sections = _split_sections(resume_text)
    skills_blob = " ".join(sections.get("skills", [])).lower()

    def cover(terms):
        if not terms: return 0.0, [], terms
        hits = []
        for t in terms:
            if (len(t) < 3 and not _is_acronym(t.upper())):
                continue
            if t in STOPWORDS_KW or t in NOISE_VERBS or t in NOISE_NOUNS:
                continue
            if t in txt or re.search(rf"\b{re.escape(t)}\b", txt):
                hits.append(t)
        matched = sorted(set(hits))
        missing = sorted(set(terms) - set(matched))
        cov = len(matched) / max(1, len(set(terms)))
        return cov, matched, missing

    cov_sk, m_sk, miss_sk = cover(jd_terms.get("skills", []))
    cov_ti, m_ti, miss_ti = cover(jd_terms.get("titles", []))
    cov_ce, m_ce, miss_ce = cover(jd_terms.get("certs", []))

    recent_txt = " ".join(" ".join(r["lines"]).lower() for r in (roles or [])[:2])

    def recency_boost(cov, matched):
        if not matched:
            return cov
        in_recent = sum(1 for t in matched if t in recent_txt)
        if in_recent:
            cov = min(1.0, cov + 0.1)
        return cov

    cov_sk = recency_boost(cov_sk, m_sk)
    cov_ti = recency_boost(cov_ti, m_ti)
    cov_ce = recency_boost(cov_ce, m_ce)

    pts_sk = int(round(15 * min(1.0, cov_sk)))
    pts_ti = int(round( 8 * min(1.0, cov_ti)))
    pts_ce = int(round( 6 * min(1.0, cov_ce)))

    # Acronym ↔ full-form coverage
    acro_ok = 0
    for a, full in jd_terms.get("acronyms", {}).items():
        if re.search(rf"\b{re.escape(a)}\b", txt) and re.search(rf"\b{re.escape(full)}\b", txt):
            acro_ok = 3
            break

    # Distribution (not only in "skills" section)
    in_skills = sum(1 for t in set(m_sk + m_ti + m_ce) if t in skills_blob)
    distribution_ok = in_skills < max(2, int(0.6 * len(set(m_sk + m_ti + m_ce))))
    pts_dist = 3 if distribution_ok else 0

    total_pts = pts_sk + pts_ti + pts_ce + acro_ok + pts_dist          # 0–35
    score_pct = int(round(100 * total_pts / 35))                        # 0–100

    matched_all = sorted(set(m_sk + m_ti + m_ce))
    missing_all = sorted(set((jd_terms.get("skills", []) or []) +
                             (jd_terms.get("titles", []) or []) +
                             (jd_terms.get("certs", [])  or [])) - set(matched_all))

    return {
        "score": score_pct,  # normalized 0–100
        "matched": matched_all,
        "missing": missing_all,
        "distribution_ok": distribution_ok,
        "components": {"skills": pts_sk, "titles": pts_ti, "certs": pts_ce, "acro": acro_ok, "distribution": pts_dist, "total_points": total_pts}
    }

def _experience_relevance(jd_text: str, roles: list, resume_text: str):
    m = re.search(r"(\d+)\+?\s+years", jd_text, re.I)
    req = int(m.group(1)) if m else None
    yoe = _years_of_experience(roles)
    years_ok = (yoe >= req) if req else (yoe >= 2)
    s_years = 6 if years_ok else max(0, int(6 * (yoe / max(1, req or 6))))

    titles = [ln.lower() for r in roles for ln in r["lines"][:2]]
    has_lead = any(re.search(r"\b(lead|manager|senior|principal)\b", t) for t in titles)
    s_senior = 4 if has_lead or yoe >= 5 else 2 if yoe >= 3 else 1

    verbs = {"implement","migrate","optimize","build","design","automate","support","troubleshoot","manage"}
    recent_txt = " ".join(" ".join(r["lines"]).lower() for r in roles[:2])
    overlap = len([v for v in verbs if v in recent_txt])
    s_recent = 3 if overlap >= 4 else 2 if overlap >= 2 else 1

    jd_verbs = {w for w in _tokenize(jd_text) if w in verbs}
    res_verbs= {w for w in _tokenize(resume_text) if w in verbs}
    s_task   = 2 if len(jd_verbs & res_verbs) >= 2 else 1 if (jd_verbs & res_verbs) else 0

    return {"score": s_years + s_senior + s_recent + s_task, "yoe": yoe, "req": req}

def _education_and_certs(resume_text: str, jd_terms: dict):
    score = 0
    if re.search(r"\b(bsc|msc|b\.?sc|m\.?sc|bachelor|master|degree)\b", resume_text, re.I):
        score += 3
    certs_found = [c for c in jd_terms.get("certs", []) if c in resume_text.lower()]
    if certs_found: score += 3
    return {"score": score, "certs_found": certs_found}

def _eligibility_location(resume_text: str):
    score = 0
    if LOC_PAT.search(resume_text): score += 1
    if re.search(r"\b(eligible to work|work authorization|right to work|visa|relocat)\b", resume_text, re.I):
        score += 2
    return {"score": score}

def _readability_brevity(resume_text: str, level: str):
    words = len(_tokenize(resume_text))
    est_pages = max(1, int(round(words/WORDS_PER_PAGE)))
    if level in ("exec","senior"):
        length_ok = est_pages <= 3
    else:
        length_ok = est_pages <= 2
    s_len = 3 if length_ok else 1

    qs = _quant_stats(resume_text)
    within = 8 <= qs["avg_words_per_bullet"] <= 20 if qs["bullet_count"] else False
    s_bul = 3 if within else 1

    dates = re.findall(r"(?:\b\d{2}/\d{4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\.?\s+\d{4})", resume_text, re.I)
    has_consistency = len(dates) >= 2
    s_con = 2 if has_consistency else 1
    return {
        "score": s_len + s_bul + s_con,
        "length_pages": est_pages,
        "bullets": qs,
        "length_component": s_len,
        "bullet_component": s_bul,
        "consistency_component": s_con,
    }

def _content_depth_penalty(resume_text: str, level: str):
    """Penalise very short / sparse resumes so 1-page thin CVs can't score high."""
    words = len(_tokenize(resume_text))
    qs = _quant_stats(resume_text)
    bullets = qs["bullet_count"]

    min_words = 300 if level in ("entry",) else 350
    min_bullets = 6 if level in ("entry",) else 8

    pts = 0; reasons = []
    if words < min_words:
        gap = min_words - words
        pts += 4 if gap < 100 else 6 if gap < 180 else 8
        reasons.append("Too little content — add more impact bullets and detail.")
    if bullets < min_bullets:
        pts += 3 if bullets >= (min_bullets - 2) else 6
        reasons.append("Too few bullet points — aim for concise 8–20 word bullets.")

    return min(12, pts), reasons

def _detect_unnecessary_sections(resume_text: str):
    found = []
    if re.search(r"(?i)^\s*references\b", resume_text, re.M): found.append("References")
    if re.search(r"(?i)^\s*objective\b",  resume_text, re.M): found.append("Objective")
    return found

def _penalties(resume_text: str, roles: list, hard_fail: bool):
    pts, reasons = 0, []
    if hard_fail: return 100, ["Image-only/unparseable PDF"]

    # missing any dates across roles
    if sum(1 for r in roles if r.get("start") or r.get("end")) == 0:
        pts += 5; reasons.append("Missing dates on roles")

    toks = _tokenize(resume_text)
    if toks:
        counts = collections.Counter(toks)
        common_actions = {"managed","led","delivered","supported","coordinated"}
        for a in common_actions: counts.pop(a, None)
        if counts:
            top, n = counts.most_common(1)[0]
            if n/len(toks) > 0.08 or n >= 40:
                pts += 4; reasons.append("Possible keyword stuffing")

    if re.search(r"\bmy journey\b", resume_text, re.I):
        pts += 3; reasons.append("Non-standard headings")

    years = sorted({r["start"] for r in roles if r.get("start")} | {r["end"] for r in roles if r.get("end")})
    for a,b in zip(years, years[1:]):
        if b - a >= 2:
            pts += 3; reasons.append("Potential gaps"); break

    if any(re.search(rf"\b{re.escape(p)}\b", resume_text, re.I) for p in PERSONAL_DATA):
        pts += 2; reasons.append("Personal data (DOB/photo/etc.)")

    return pts, reasons

@resumes_bp.post("/build-resume")
@login_required
def build_resume():
    data  = request.get_json(force=True) or {}
    theme = (data.get("theme") or "modern").lower()
    fmt   = (data.get("format") or "html").lower()

    ctx = _normalize_ctx(data)

    # 1) Paid/watermark flags for the template (explicit)
    plan = (getattr(current_user, "plan", "free") or "free").lower()
    is_paid = plan in ("standard", "premium")
    wm_text = "" if is_paid else "JOBCUS.COM"

    tpl_ctx = {
        "name":           ctx.get("name", ""),
        "title":          ctx.get("title", ""),
        "contact":        ctx.get("contact", ""),
        "summary":        ctx.get("summary", ""),
        "skills":         ctx.get("skills", []),
        "links":          ctx.get("links", []),
        "experience":     ctx.get("experience", []),
        "education":      ctx.get("education", []),
        "certifications": ctx.get("certifications", []),
        "projects_awards": ctx.get("projects_awards", []),   # ← add this
        "is_paid": is_paid,
        "wm_text": wm_text,
        "for_pdf": (fmt == "pdf"),
    }

    template_path = f"resumes/{'minimal' if theme == 'minimal' else 'modern'}.html"
    html = render_template(template_path, **tpl_ctx)

    # 2) Optional PDF hardening: if template forgot watermark, add a light overlay
    if fmt == "pdf" and not is_paid:
        if "data-watermark" not in html:
            html = html.replace("<body", f'<body data-watermark="{wm_text}"', 1)
        if "data-watermark]::before" not in html:
            wm_css = """
<style>
  body[data-watermark]::before{
    content: attr(data-watermark);
    position: fixed;
    top: 50%; left: 50%;
    transform: translate(-50%,-50%) rotate(-32deg);
    font: 72px/1.1 Arial, Helvetica, sans-serif;
    color: rgba(0,0,0,.12);
    letter-spacing: 2px;
    white-space: nowrap;
    user-select: none;
    pointer-events: none;
    z-index: 9999;
  }
</style>"""
            html = html.replace("</head>", wm_css + "</head>", 1)

    if fmt == "pdf":
        if not feature_enabled(plan, "downloads"):
            return jsonify(
                error="upgrade_required",
                message="File downloads are available on Standard and Premium."
            ), 403

        stylesheets = [CSS(string=PDF_CSS_OVERRIDES)]
        pdf_css_path = os.path.join(current_app.root_path, "static", "pdf.css")
        if os.path.exists(pdf_css_path):
            stylesheets.append(CSS(filename=pdf_css_path))

        pdf_bytes = HTML(string=html, base_url=current_app.root_path).write_pdf(stylesheets=stylesheets)
        resp = make_response(pdf_bytes)
        resp.headers["Content-Type"] = "application/pdf"
        resp.headers["Content-Disposition"] = "inline; filename=resume.pdf"
        return resp

    resp = make_response(html)
    resp.headers["Content-Type"] = "text/html; charset=utf-8"
    return resp

# ---------- Template-based resume (DOCX) ----------
@resumes_bp.post("/build-resume-docx")
def build_resume_docx():
    try:
        data = request.get_json(force=True) or {}

        # helpers to normalize inputs (strings or arrays)
        def _lines(val):
            if isinstance(val, list):
                return [str(x).strip() for x in val if str(x).strip()]
            if isinstance(val, str):
                parts = [p.strip() for p in val.replace("•", "").split("\n")]
                if len(parts) == 1 and "," in parts[0]:
                    parts = [p.strip() for p in parts[0].split(",")]
                return [p for p in parts if p]
            return []

        def _sections(val):
            if isinstance(val, list):
                return val
            if isinstance(val, str) and val.strip():
                return [{"bullets": [v.strip() for v in val.split("\n") if v.strip()]}]
            return []

        full_name = data.get("name") or data.get("fullName") or data.get("full_name") or ""
        title     = data.get("title", "")
        contact   = data.get("contact", "")
        summary   = data.get("summary", "")
        skills    = _lines(data.get("skills", ""))
        certs     = _lines(data.get("certifications", ""))
        experience = _sections(data.get("experience", []))
        education  = _sections(data.get("education", []))
        projects_awards = _lines(data.get("projects_awards", "") or data.get("projectsAndAwards", ""))

        doc = Document()

        # Header
        if full_name: doc.add_heading(full_name, level=0)
        if title:     doc.add_paragraph(title)
        if contact:   doc.add_paragraph(contact)

        # Summary
        if summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(summary)

        # Skills
        if skills:
            doc.add_heading("Skills", level=1)
            for s in skills:
                doc.add_paragraph(s, style="List Bullet")

        # Experience
        if experience:
            doc.add_heading("Relevant Experience", level=1)
            for e in experience:
                role    = (e.get("role") or "").strip()
                company = (e.get("company") or "").strip()
                header  = " – ".join([p for p in [role, company] if p])
                if header:
                    p = doc.add_paragraph(header)
                    try:
                        for run in p.runs: run.bold = True
                    except:
                        pass

                location = (e.get("location") or "").strip()
                if location:
                    doc.add_paragraph(location)

                start = (e.get("start") or "").strip()
                end   = (e.get("end")   or "").strip()
                if start or end:
                    doc.add_paragraph(" · ".join([p for p in [start, end] if p]))

                for b in (e.get("bullets") or []):
                    if str(b).strip():
                        doc.add_paragraph(str(b).strip(), style="List Bullet")

        # Education
        if education:
            doc.add_heading("Education", level=1)
            for ed in education:
                # ---- Title line (degree/program) ----
                degree = (ed.get("degree") or ed.get("program") or ed.get("title") or "").strip()
                if degree:
                    p = doc.add_paragraph()
                    run = p.add_run(degree)
                    try:
                        run.bold = True
                    except Exception:
                        pass
        
                # ---- Sub line: Institution | Location – Start – End ----
                school_line = " | ".join([p for p in [
                    (ed.get("school") or "").strip(),
                    (ed.get("location") or "").strip()
                ] if p])
        
                start = (ed.get("graduatedStart") or ed.get("start") or "").strip()
                end   = (ed.get("graduated")      or ed.get("end")   or "").strip()
        
                date_bits = [x for x in [start, end] if x]
                date_line = " – ".join(date_bits)   # en dash
        
                line = f"{school_line}{(' – ' + date_line) if date_line else ''}"
                if line:
                    doc.add_paragraph(line)

        # Certifications
        if certs:
            doc.add_heading("Other Education & Certifications", level=1)
            for c in certs:
                doc.add_paragraph(c, style="List Bullet")

        # Projects & Awards
        if projects_awards:
            doc.add_heading("Projects & Awards", level=1)
            for item in projects_awards:
                doc.add_paragraph(item, style="List Bullet")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        return send_file(
            buf,
            as_attachment=True,
            download_name="resume.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        # requires `from flask import current_app`
        current_app.logger.exception("DOCX build failed")
        return jsonify(error="docx_failed", message=str(e)), 500

    
# ---------- 4) AI resume optimisation ----------
@resumes_bp.route("/api/optimize-resume", methods=["POST"])
@login_required
def optimize_resume():
    plan = (getattr(current_user, "plan", "free") or "free").lower()
    if not feature_enabled(plan, "optimize_ai"):
        return jsonify(
            error="upgrade_required",
            message="Optimize with AI is available on Standard and Premium."
        ), 403

    client = current_app.config["OPENAI_CLIENT"]
    data = request.get_json(force=True)
    resume_text = ""

    if data.get("pdf"):
        try:
            pdf_bytes = base64.b64decode(data["pdf"])
            reader = PdfReader(BytesIO(pdf_bytes))
            resume_text = "\n".join((p.extract_text() or "") for p in reader.pages)
            if not resume_text.strip():
                return jsonify(error="PDF content appears to have no selectable text (likely scanned). Upload a text-based PDF or DOCX."), 400
        except Exception:
            logging.exception("PDF Decode Error")
            return jsonify(error="Unable to extract PDF text (corrupt or scanned). Upload a text-based PDF or DOCX."), 400
    elif data.get("docx"):
        try:
            docx_bytes = base64.b64decode(data["docx"])
            d = docx.Document(BytesIO(docx_bytes))
            resume_text = "\n".join(p.text for p in d.paragraphs)
            if not resume_text.strip():
                return jsonify(error="DOCX appears empty. Please check the file content."), 400
        except Exception:
            logging.exception("DOCX Decode Error")
            return jsonify(error="Unable to read DOCX. Re-save as .docx and try again."), 400
    elif data.get("text"):
        resume_text = (data.get("text") or "").strip()
        if not resume_text:
            return jsonify(error="No text provided"), 400
    else:
        return jsonify(error="No resume data provided"), 400

    prompt = (
        "You are an expert ATS resume optimizer. Rewrite the following resume in plain text, "
        "using strong action verbs, consistent bullets, relevant keywords, and fixing grammar/repetition.\n\n"
        f"Original resume:\n\n{resume_text}"
    )
    try:
        model = current_app.config.get("CHOOSE_MODEL", lambda r: "gpt-4o-mini")(None)
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role":"user","content":prompt}],
            temperature=0.3
        )
        optimized = (resp.choices[0].message.content or "").strip()

        # Log usage (optimize resume LLM pass)
        log_ai_usage(
            feature="optimize_resume",
            model="gpt-4o",
            resp=resp,
            extra={
                "status": "ok",
                "cost_gbp": estimate_cost_gbp(
                    "gpt-4o",
                    getattr(resp.usage, "prompt_tokens", None),
                    getattr(resp.usage, "completion_tokens", None),
                )
            }
        )

        optimized = re.sub(r"```(?:[\s\S]*?)```", "", optimized).strip()
        return jsonify({"optimized": optimized})
    except Exception:
        logging.exception("Resume optimization error")
        return jsonify(error="Resume optimization failed"), 500

# ---------- 5) AI generate resume (JSON for template) ----------
@resumes_bp.post("/generate-resume")
def generate_resume():
    client = current_app.config["OPENAI_CLIENT"]
    data = request.get_json(force=True) or {}

    prompt = f"""
Return ONLY valid JSON (no backticks) with this exact schema:

{{
  "name": "...",
  "title": "...",
  "contact": "...",
  "summary": "...",
  "skills": ["..."],
  "links": [{{"url":"...", "label":""}}],
  "experience": [
    {{"role":"...", "company":"...", "location":"", "start":"", "end":"", "bullets":["..."]}}
  ],
  "education": [
    {{"degree":"...", "school":"...", "location":"", "graduated":""}}
  ],
  "certifications": ["..."]
}}

User input:
fullName: {data.get('fullName',"")}
firstName: {data.get('firstName',"")}
lastName: {data.get('lastName',"")}
title: {data.get('title',"")}
contact: {data.get('contact',"")}
summary: {data.get('summary',"")}
education (free text): {data.get('education',"")}
experience (free text): {data.get('experience',"")}
skills (free text): {data.get('skills',"")}
certifications (free text): {data.get('certifications',"")}
portfolio: {data.get('portfolio',"")}
""".strip()

    def _coerce_list(val):
        if isinstance(val, list): return val
        if not val: return []
        return [s.strip() for s in str(val).replace("\r", "").split("\n") if s.strip()]

    try:
        model = current_app.config.get("CHOOSE_MODEL", lambda r: "gpt-4o-mini")(None)
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role":"user","content":prompt}],
            temperature=0.2
        )
        content = resp.choices[0].message.content.strip()

        # Log usage (resume builder LLM pass)
        log_ai_usage(
            feature="resume_builder",
            model="gpt-4o",
            resp=resp,
            extra={
                "status": "ok",
                "cost_gbp": estimate_cost_gbp(
                    "gpt-4o",
                    getattr(resp.usage, "prompt_tokens", None),
                    getattr(resp.usage, "completion_tokens", None),
                )
            }
        )
        content = re.sub(r"```(?:json)?", "", content).strip()
        ctx = json.loads(content)

        first = (data.get("firstName") or "").strip()
        last  = (data.get("lastName") or "").strip()
        derived = " ".join([p for p in (first, last) if p]).strip()
        ctx.setdefault("name",  (data.get("fullName") or derived or ""))
        ctx.setdefault("title", (data.get("title") or ""))
        ctx.setdefault("contact", (data.get("contact") or ""))

        if isinstance(ctx.get("skills"), str):
            ctx["skills"] = [s.strip() for s in ctx["skills"].replace(",", "\n").splitlines() if s.strip()]
        ctx["certifications"] = _coerce_list(ctx.get("certifications")) or _coerce_list(data.get("certifications"))

        return jsonify(context=ctx, aiUsed=True)

    except RateLimitError:
        current_app.logger.error("OpenAI quota/429 in /generate-resume")
        return jsonify(context=naive_context(data), aiUsed=False, error_code="quota_or_error")

    except Exception:
        logging.exception("Generation failed")
        return jsonify(context=naive_context(data), aiUsed=False, error_code="error")

@resumes_bp.route("/ai/suggest", methods=["POST"])
def ai_suggest():
    data  = request.get_json(force=True) or {}
    field = (data.get("field") or "general").strip().lower()
    ctx   = data.get("context") or {}
    client = current_app.config.get("OPENAI_CLIENT")

    def normalize(text="", items=None):
        items = items if isinstance(items, list) else None
        if not items and text:
            items = [ln.strip("•- ").strip() for ln in text.splitlines() if ln.strip()]
        if not text and items:
            text = "\n".join(items)
        return {"text": text or "", "list": items or [], "suggestions": items or []}

    # ---- Cover letter (letter-style body, richer content) ----
    if field in ("coverletter", "coverletter_from_analyzer", "cover_letter"):
        name   = (ctx.get("name") or "").strip()
        title  = (ctx.get("title") or "professional").strip()
        cl     = ctx.get("coverLetter") or ctx or {}
        company = (cl.get("company") or "your company").strip()
        role    = (cl.get("role") or "the role").strip()
        manager = (cl.get("manager") or "Hiring Manager").strip()
        tone    = (cl.get("tone") or "professional").strip()

        jd   = (ctx.get("jd") or "").strip()
        rsum = (ctx.get("resumeText") or "").strip()

        prompt = (
            "Write a UK-English cover letter BODY (no greeting and no closing) for a job application.\n"
            f"Role: {role}\nCompany: {company}\nCandidate: {name or 'the candidate'} ({title})\n"
            f"Tone: {tone} but warm and confident. Length: ~220–320 words in 3–5 short paragraphs.\n"
            "Prioritise clear impact, quantified results, collaboration, tooling, and domain fit.\n"
            "Use concise sentences, avoid clichés and buzzwords, and do not include lists or bullets.\n"
            "Return PLAIN TEXT only — DO NOT add 'Dear ...' or 'Yours sincerely'.\n\n"
            f"Resume excerpt (optional):\n{rsum[:2000]}\n\n"
            f"Job description (optional):\n{jd[:2000]}"
        )

        if client:
            try:
                model = current_app.config.get("CHOOSE_MODEL", lambda r: "gpt-4o-mini")(None)
                resp = client.chat.completions.create(
                    model=model,
                    messages=[{"role":"user","content":prompt}],
                    temperature=0.6,
                    max_tokens=600,
                )
                out = (resp.choices[0].message.content or "").strip()
                out = re.sub(r"```(?:\w+)?", "", out).strip()
                return jsonify({"text": out, "list": [], "suggestions": []})
            except Exception as e:
                current_app.logger.warning("cover letter AI error; using fallback: %s", e)

        text = (
            f"As an experienced {title.lower()} with a strong record of supporting teams and customers, "
            f"I’m excited to apply for the {role} role at {company}. I combine hands-on technical ability with "
            f"a pragmatic, service-oriented approach that focuses on reliability, responsiveness, and measurable outcomes.\n\n"
            "In previous roles I resolved complex issues across hardware, software, and networks while maintaining high CSAT. "
            "I introduced lightweight runbooks to speed up triage, reduced repeat incidents through root-cause fixes, "
            "and partnered with cross-functional teams to tighten feedback loops. Notable wins include cutting average resolution time "
            "by over 25% and driving onboarding improvements that reduced first-week tickets by double digits.\n\n"
            "Beyond troubleshooting, I care deeply about clear communication. I translate technical details into plain language, "
            "set expectations transparently, and prioritise thoughtfully when demand spikes. I’m comfortable owning a queue, "
            "coordinating escalations, and documenting knowledge so solutions scale.\n\n"
            "I’m confident I can bring the same reliability and customer focus to your team and help deliver fast, friendly, "
            "and secure support at scale."
        )
        return jsonify({"text": text, "list": [], "suggestions": []})


    # ---- Compact context for resume prompts ----
    def compact_context(c):
        parts = []
        nm  = c.get("name");   ti = c.get("title")
        sm  = c.get("summary"); ct = c.get("contact")
        if nm or ti: parts.append(f"Name/Title: {nm or ''} {ti or ''}".strip())
        if ct: parts.append(f"Contact: {ct}")
        exps = c.get("experience") or []
        if exps:
            e0 = exps[0]
            parts.append(
                f"Recent role: {e0.get('role','')} at {e0.get('company','')} "
                f"({e0.get('location','')}) {e0.get('start','')}–{e0.get('end','')}"
            )
        if sm: parts.append(f"Existing summary: {sm}")
        return "\n".join(parts)[:2000]

    base_ctx = compact_context(ctx)

    if field == "summary":
        prompt = (
            "Write a single, crisp professional summary (2–3 sentences) for a resume. "
            "Emphasize years of experience, role scope, and 1–2 measurable achievements. "
            "No bullet points. Be concise and ATS-friendly.\n\n"
            f"Context:\n{base_ctx}"
        )
    elif field == "highlights":
        idx = data.get("index")
        exp = None
        if isinstance(idx, int):
            exps = ctx.get("experience") or []
            if 0 <= idx < len(exps):
                exp = exps[idx]
        exp_ctx = ""
        if exp:
            exp_ctx = (
                f"Role: {exp.get('role','')}\n"
                f"Company: {exp.get('company','')}\n"
                f"Location: {exp.get('location','')}\n"
                f"Dates: {exp.get('start','')} – {exp.get('end','')}\n"
                f"Existing bullets: {(exp.get('bullets') or [])}"
            )
        prompt = (
            "Write exactly 3–5 strong resume bullet points for this role. "
            "Return plain lines (no numbering). Start with an action verb; keep each under ~22 words; quantify impact.\n\n"
            f"Global context:\n{base_ctx}\n\nJob context:\n{exp_ctx}"
        )
    else:
        prompt = (
            "Write 3 concise, outcomes-focused resume bullet points. "
            "Return plain lines (no numbering). Use action verbs and metrics.\n\n"
            f"Context:\n{base_ctx}"
        )

    if not client:
        if field == "summary":
            return jsonify(normalize(
                "Results-driven professional with experience delivering measurable impact across X, Y, and Z. "
                "Recognized for A, B, and C; partners cross-functionally to ship results."
            ))
        return jsonify(normalize(items=[
            "Increased X by Y% by doing Z",
            "Reduced A by B% through C initiative",
            "Led D cross-functional effort resulting in E"
        ]))

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert resume writer. Be concise and ATS-friendly."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5,
            max_tokens=220,
        )
        out = (resp.choices[0].message.content or "").strip()
        return jsonify(normalize(out))
    except Exception as e:
        current_app.logger.warning("ai_suggest error; returning fallback: %s", e)
        if field == "summary":
            return jsonify(normalize(
                "Results-driven professional with experience delivering measurable impact across X, Y, and Z. "
                "Recognized for A, B, and C; partners cross-functionally to ship results."
            ))
        return jsonify(normalize(items=[
            "Increased X by Y% by doing Z",
            "Reduced A by B% through C initiative",
            "Led D cross-functional effort resulting in E"
        ]))

# ---------- Cover Letter builder ----------
@resumes_bp.route("/build-cover-letter", methods=["POST"])
@login_required
def build_cover_letter():
    """
    Preview (HTML) is always allowed.
    Only gate the actual download (PDF) like the resume builder does.
    Always normalize the letter body so templates can use {{ cover_body }}.
    Fail gracefully (JSON) instead of throwing a 500 HTML page.
    """
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        current_app.logger.exception("build-cover-letter: bad JSON")
        return jsonify(error="invalid_request", message="Invalid JSON body"), 400

    fmt         = (data.get("format") or "html").lower()
    letter_only = str(data.get("letter_only")).lower() in ("1", "true", "yes", "on")

    # --- Gate ONLY file download (PDF) ---
    if fmt == "pdf":
        plan = (getattr(current_user, "plan", "free") or "free").lower()
        if not feature_enabled(plan, "downloads"):
            pricing_url = "/pricing"
            return jsonify(
                error="upgrade_required",
                message="File downloads are available on Standard and Premium.",
                message_html=f'File downloads are available on Standard and Premium. <a href="{pricing_url}">Upgrade now →</a>',
                pricing_url=pricing_url
            ), 403

    # --- Normalize context and body ---
    sender    = data.get("sender")    or {}
    recipient = data.get("recipient") or {}
    cl        = data.get("coverLetter") or {}
    cover_body = (
        data.get("cover_body") or
        data.get("draft") or
        data.get("body") or
        cl.get("draft") or
        cl.get("text") or
        ""
    )

    tpl_ctx = {
        # top-level fields (optional)
        "name":   data.get("name")  or data.get("fullName") or "",
        "title":  data.get("title") or "",
        "contact": data.get("contact") or "",
        # structured blocks
        "sender": sender,
        "recipient": recipient,
        # always provide both names to avoid template mismatches
        "cover_body": cover_body,
        "draft": cover_body,
        # flags some templates expect
        "letter_only": letter_only,
        "for_pdf": (fmt == "pdf"),
    }

    # --- Choose template robustly & render safely ---
    candidates = []
    if letter_only:
        candidates = ["cover-letter-print.html", "cover-letter.html"]
    else:
        candidates = ["cover-letter.html", "cover-letter-print.html"]

    html = None
    for tpl in candidates:
        try:
            html = render_template(tpl, **tpl_ctx)
            break
        except TemplateNotFound:
            continue
        except Exception as e:
            current_app.logger.exception("build-cover-letter: render error in %s", tpl)
            return jsonify(error="template_error", message=str(e)), 500

    if html is None:
        current_app.logger.error("build-cover-letter: no template found (tried %s)", candidates)
        return jsonify(error="template_not_found", message="Cover letter template missing"), 500

    # --- Return as PDF or HTML ---
    if fmt == "pdf":
        try:
            pdf_bytes = HTML(string=html, base_url=current_app.root_path)\
                .write_pdf(stylesheets=[CSS(string="@page{size:A4;margin:0.75in}")])
            resp = make_response(pdf_bytes)
            resp.headers["Content-Type"] = "application/pdf"
            resp.headers["Content-Disposition"] = "inline; filename=cover-letter.pdf"
            return resp
        except Exception as e:
            current_app.logger.exception("build-cover-letter: PDF generation failed")
            return jsonify(error="pdf_failed", message=str(e) or "PDF generation failed"), 500

    # default: HTML preview
    resp = make_response(html)
    resp.headers["Content-Type"] = "text/html; charset=utf-8"
    return resp
    
# ----------- AI Helper -----------------
@resumes_bp.post("/api/ai-helper")
@api_login_required
def ai_helper():
    client = current_app.config.get("OPENAI_CLIENT")
    try:
        data = request.get_json(force=True) or {}
    except Exception:
        return jsonify(error="bad_request", message="Invalid JSON"), 400

    field = (data.get("field") or "").strip()
    ctx   = data.get("context") or {}
    idx   = int(data.get("index") or 0)

    def chat(msg, temp=0.4, max_tokens=600):
        if not client:
            return ""
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"user","content":msg}],
            temperature=temp,
            max_tokens=max_tokens
        )
        return (r.choices[0].message.content or "").strip()

    if field == "resume_analysis":
        bullets = ctx.get("bullets") or []
        summary = (ctx.get("summary") or "").strip()
        prompt = f"""You analyze resume bullets. For each bullet, list concrete problems and a stronger 1-line rewrite.
Return strict JSON: {{"issues":[{{"text":"...", "problems":["..."], "suggestion":"..."}}]}}
Bullets:\n- """ + "\n- ".join(bullets[:20]) + f"\n\nSummary: {summary or 'N/A'}"
        raw = chat(prompt, temp=0.0)
        try:
            import re, json
            block = re.search(r"\{[\s\S]*\}\s*$", raw).group(0)
            parsed = json.loads(block)
            return jsonify(issues=parsed.get("issues", []))
        except Exception:
            return jsonify(issues=[])
    if field == "rewrite_bullet":
        text = (ctx.get("text") or "").strip()
        prompt = f"Rewrite this resume bullet (1 line, metric/outcome, strong verb). Return only the rewritten line.\n{text}"
        return jsonify(text=chat(prompt, temp=0.4, max_tokens=140))
    if field == "highlights":
        role = (ctx.get("experience", [{}])[idx].get("role") if ctx.get("experience") else "") or (ctx.get("title") or "professional")
        prompt = f"Write 4 concise resume bullets for a {role}. Each starts with a strong verb and includes a metric. Split by newlines."
        lines = [s.strip("• ").strip() for s in chat(prompt).splitlines() if s.strip()]
        return jsonify(list=lines[:6])
    if field == "summary":
        title = ctx.get("title") or "professional"
        skills = ", ".join(ctx.get("skills") or [])[:120] or "relevant strengths"
        return jsonify(text=chat(f"2–3 sentence resume summary for a {title} highlighting {skills}.", temp=0.5, max_tokens=180))
    return jsonify(error="unknown_field", message="Unsupported field"), 400

# ---------- 3) AI resume analysis ----------
@resumes_bp.route("/api/resume-analysis", methods=["POST"])
@api_login_required
def resume_analysis():
    supabase_admin = current_app.config["SUPABASE_ADMIN"]
    plan = (getattr(current_user, "plan", "free") or "free").lower()

    # NEW: device/user-scoped free-plan guard (no router IP)
    ok, payload = allow_free_use(current_user.id, plan)
    if not ok:
        return jsonify({
            "error": "too_many_free_accounts",
            "message": payload.get("message") or "Free usage limit reached for this device."
        }), 429

    # FIX: align quota key with limits/UI: resume_analyzer
    allowed, info = check_and_increment(supabase_admin, current_user.id, plan, "resume_analyzer")
    if not allowed:
        PRICING_URL = "https://www.jobcus.com/pricing"
        info.setdefault("error", "quota_exceeded")
        info.setdefault("message", "You’ve reached your plan limit for this feature.")
        info.setdefault("message_html", f'You’ve reached your plan limit for this feature. <a href="{PRICING_URL}">Upgrade now →</a>')
        info.setdefault("pricing_url", PRICING_URL)
        return jsonify(info), 402

    # Additional hour/day call caps
    okh, infh = check_and_increment(supabase_admin, current_user.id, plan, "resume_analyzer_hour")
    if not okh:
        return jsonify(infh), 429
    
    okd, infd = check_and_increment(supabase_admin, current_user.id, plan, "resume_analyzer_day")
    if not okd:
        return jsonify(infd), 429

    """
    ATS scoring model (100 pts) + penalties, aligned to industry best practice.
    Preserves your existing response shape for the dashboard.
    """
    client = current_app.config.get("OPENAI_CLIENT")
    data = request.get_json(force=True) or {}

    # ---- Decode input ----
    resume_text, file_kind, raw_bytes = "", None, None
    try:
        if data.get("pdf"):
            raw_bytes = base64.b64decode(data["pdf"])
            reader = PdfReader(BytesIO(raw_bytes))
            resume_text = "\n".join((p.extract_text() or "") for p in reader.pages)
            resume_text = _normalize_extracted_text(resume_text)  # normalize
            file_kind = "pdf"
    
        elif data.get("docx"):
            raw_bytes = base64.b64decode(data["docx"])
            buf = BytesIO(raw_bytes)
            d = docx.Document(buf)
            resume_text = "\n".join(p.text for p in d.paragraphs)
            resume_text = _normalize_extracted_text(resume_text)  # normalize
            file_kind = "docx"
    
        elif data.get("text"):
            resume_text = (data["text"] or "").strip()
            resume_text = _normalize_extracted_text(resume_text)  # normalize
            file_kind = "text"
    
        else:
            return jsonify(error="No resume data provided"), 400
    
    except Exception:
        current_app.logger.exception("resume-analysis: decode failed")
        return jsonify(error="Could not read the resume file"), 400
    
    if not (resume_text or "").strip():
        return jsonify(error="Could not extract any text"), 400
    
    job_desc = (data.get("jobDescription") or "").strip()
    job_role = (data.get("jobRole") or "").strip()
    level    = (data.get("careerLevel") or "mid").lower()

    # ---- Deterministic ATS categories ----
    pf_score, hard_fail = _parse_and_format_score(resume_text, file_kind, raw_bytes)  # /10
    sec_struct   = _sections_structure_score(resume_text)                              # /15
    roles        = _extract_roles_with_dates(resume_text)

    # Build JD terms with fallbacks: JD > role > summary/skills
    jd_source = (job_desc or "").strip()
    if not jd_source and job_role:
        jd_source = job_role
    if not jd_source:
        secs = _split_sections(resume_text)
        summary_blob = " ".join(secs.get("summary", []))
        skills_blob  = " ".join(secs.get("skills", []))
        jd_source = f"{summary_blob}\n{skills_blob}".strip()

    jd_terms = _build_jd_terms(jd_source) if jd_source else {"all": [], "skills": [], "titles": [], "certs": [], "acronyms": {}}

    kw_match     = _keyword_match_to_jd(resume_text, jd_terms, roles)   # points 0–35
    kw_points    = kw_match["score"]

    # Clean the keyword lists before returning
    matched = _clean_kw_list(kw_match.get("matched", []))
    missing = _clean_kw_list(kw_match.get("missing", []))

    exp_rel      = _experience_relevance(job_desc, roles, resume_text)  # /15
    quant        = _quant_stats(resume_text)                             # achievements
    ach_score    = (5 if quant["pct_with_numbers"] >= 50 else int(5 * quant["pct_with_numbers"]/50)) \
                   + (3 if quant["pct_action_starts"] >= 70 else int(3 * quant["pct_action_starts"]/70))  # /8
    edu_certs    = _education_and_certs(resume_text, jd_terms)           # /6
    elig_loc     = _eligibility_location(resume_text)                    # /3
    read_brief   = _readability_brevity(resume_text, level)              # /8

    # ---- Penalties ----
    depth_pts, depth_reasons = _content_depth_penalty(resume_text, level)
    pen_pts, pen_reasons = _penalties(resume_text, roles, hard_fail)
    pen_pts += depth_pts
    pen_reasons += depth_reasons
    if not sec_struct["std"]["experience"]:
        pen_pts += 5; pen_reasons.append("Missing Work Experience section")
    if not sec_struct["std"]["education"]:
        pen_pts += 3; pen_reasons.append("Missing Education section")

    # ---- Raw score (0–100) ----
    score = 0
    score += pf_score                              # 10
    score += min(15, sec_struct["score"])          # 15
    score += kw_points                             # 35
    score += min(15, exp_rel["score"])             # 15
    score += min(8,  ach_score)                    # 8
    score += min(6,  edu_certs["score"])           # 6
    score += min(3,  elig_loc["score"])            # 3
    score += min(8,  read_brief["score"])          # 8

    score = max(0, min(100, score - min(20, pen_pts)))

    # ---- Gates/caps so thin resumes can't look “good” ----
    cap = 100
    if job_desc:
        if kw_points < 14 or not kw_match["distribution_ok"]:
            cap = min(cap, 69)  # Needs revision
    if sec_struct["score"] < 9:
        cap = min(cap, 69)
    if quant["bullet_count"] < 6:
        cap = min(cap, 74)
    score = min(score, cap)

    # ---- Optional LLM pass for qualitative commentary ----
    llm = {
        "analysis": {"issues": [], "strengths": []},
        "suggestions": [],
        "writing": {"readability": "", "repetition": [], "grammar": []},
        "relevance": {"role": job_role, "score": 0, "explanation": "", "aligned_keywords": [], "missing_keywords": []},
    }
    
    if current_app.config.get("OPENAI_CLIENT"):
        try:
            client = current_app.config["OPENAI_CLIENT"]
            role_or_desc = (
                f"Job description:\n{job_desc}"
                if job_desc
                else (f"Target role: {job_role}" if job_role else "No job description provided.")
            )
            prompt = f"""
You are an ATS-certified resume analyst. Return ONLY valid JSON (no backticks) with:
{{
  "analysis": {{"issues": ["..."], "strengths": ["..."]}},
  "suggestions": ["..."],
  "writing": {{
    "readability": "Grade level (e.g., Grade 8–10 / B2)",
    "repetition": [{{"term":"managed","count":5,"alternatives":["led","owned","coordinated"]}}],
    "grammar": ["Short, actionable fixes."]
  }},
  "relevance": {{
    "role": "{job_role}",
    "score": 0-100,
    "explanation": "1–2 sentences",
    "aligned_keywords": ["..."],
    "missing_keywords": ["..."]
  }}
}}
Context:
{role_or_desc}

Resume:
{resume_text[:8000]}
""".strip()
    
            resp = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0,
            )
            content = (resp.choices[0].message.content or "").strip()
    
            # (usage logging + JSON extraction)
            content = re.sub(r"```(?:json)?", "", content)
            s, e = content.find("{"), content.rfind("}")
            if s >= 0 and e > s:
                js = json.loads(content[s:e + 1])
                for k, v in js.items():
                    llm[k] = v
    
        except Exception:
            current_app.logger.warning(
                "LLM step skipped; continuing with deterministic output", exc_info=True
            )
        
    # ---- Map to your UI breakdown bars ----
    breakdown = {
        "formatting": int(round((pf_score/10)*100)),
        "keywords":   kw_match["score"],                # 0–100 already
        "sections":   int(round((sec_struct["score"]/15)*100)),
        "readability": int(round((read_brief["score"]/8)*100)),
        "length":     int(round((read_brief.get("length_component", 0)/3)*100))
                      if isinstance(read_brief.get("length_component"), int)
                      else int(round((min(3, read_brief["score"])/3)*100)),
        "parseable":  (not hard_fail),
    }

    # ----- CAP headline (match what you show on the UI) -----
    visible = [
        int(breakdown["formatting"]),
        int(breakdown["sections"]),
        int(breakdown.get("keywords", 0)),
        int(breakdown["readability"]),
        int(breakdown["length"]),
        100 if breakdown["parseable"] else 0
    ]
    avg_visible = int(round(sum(visible) / len(visible)))
    all_hundred = all(v >= 100 for v in visible)
    headline = score if all_hundred else min(score, avg_visible)

    # ----- Diagnostics (build BEFORE using) -----
    diagnostics = {
        "achievements": quant,
        "eligibility_location": elig_loc,
        "education_certs": edu_certs,
        "experience_relevance": exp_rel,
        "keyword_distribution_ok": kw_match["distribution_ok"],
        "keyword_components": kw_match["components"],
        "penalties": {"points": min(20, pen_pts), "reasons": pen_reasons},
        "roles_parsed": roles[:4],
        "length_pages_est": read_brief["length_pages"]
    }

    # helper if you don't already have one
    def _is_acronym(s: str) -> bool:
        s = (s or "").strip()
        return len(s) >= 2 and len(s) <= 6 and s.isupper()
    
    # Clean up fragmenty repetition terms from LLM output
    rep = (llm.get("writing") or {}).get("repetition") or []
    clean_rep = []
    for item in rep:
        term = (item.get("term") or "").strip()
        if not term:
            continue
        # reject 1–2 letter noise unless it's a legit acronym
        if len(term) < 3 and not _is_acronym(term):
            continue
        # drop common truncated stems that come from broken extraction
        if re.search(r"(?:projec|qualit|requir|deliv|suppor|managemen|operat|depar|technolog)$", term.lower()):
            continue
        clean_rep.append(item)
    
    llm.setdefault("writing", {})["repetition"] = clean_rep

    # ----- Single, final response payload -----
    out = {
        "score": int(headline),
        "analysis": {
            "issues": llm.get("analysis", {}).get("issues", []),
            "strengths": llm.get("analysis", {}).get("strengths", []),
        },
        "suggestions": llm.get("suggestions", []),
        "breakdown": breakdown,
        "keywords": {"matched": matched, "missing": missing},
        "sections": {
            "present": [k for k, v in sec_struct["std"].items() if v],
            "missing": [k for k, v in sec_struct["std"].items() if not v],
        },
        "writing": {
            "readability": llm.get("writing", {}).get("readability", ""),
            "repetition": llm.get("writing", {}).get("repetition", []),
            "grammar": llm.get("writing", {}).get("grammar", []),
        },
        "relevance": llm.get("relevance", {}),
        "diagnostics": diagnostics,
    }

    # --- Deterministic fallback suggestions (never leave empty) ---
    if not out["suggestions"]:
        fallback = []

        # 1) Keywords coverage
        if missing:
            sample = ", ".join(missing[:8])
            fallback.append(
                f"Add role-specific keywords naturally (only if true). Consider: {sample}."
            )
        if not kw_match.get("distribution_ok", True) or kw_points < 20:
            fallback.append(
                "Spread key skills across summary, recent roles, and skills section to improve ATS term distribution."
            )

        # 2) Sections / structure
        if not sec_struct["std"]["experience"]:
            fallback.append(
                "Add a Work Experience section with employer, title, location, and Month YYYY dates."
            )
        if not sec_struct["std"]["education"]:
            fallback.append(
                "Add an Education section (degree, school, graduation year)."
            )
        if not sec_struct["std"].get("summary"):
            fallback.append(
                "Add a 2–3 sentence Professional Summary tailored to the target role."
            )
        if not sec_struct["reverse_chrono"]:
            fallback.append(
                "Reorder roles into reverse-chronological order (most recent first)."
            )

        # 3) Achievements & bullets
        if quant["bullet_count"] < 6:
            fallback.append(
                "Increase bullet count; aim for 3–5 bullets per recent role."
            )
        if quant["pct_with_numbers"] < 40:
            fallback.append(
                "Quantify impact (%, #, £/$). E.g., “Reduced resolution time by 35%”."
            )
        if quant["avg_words_per_bullet"] and (
            quant["avg_words_per_bullet"] < 8 or quant["avg_words_per_bullet"] > 20
        ):
            fallback.append(
                "Keep bullets concise (8–20 words) focused on outcomes, not duties."
            )
        if quant["pct_action_starts"] < 60:
            fallback.append(
                "Start bullets with strong action verbs (Led, Built, Reduced, Automated…)."
            )

        # 4) Readability / length
        if breakdown["length"] < 70:
            fallback.append(
                f"Adjust length to 1–2 pages (estimated {read_brief['length_pages']} page(s))."
            )
        if breakdown["readability"] < 60:
            fallback.append(
                "Simplify wording and shorten long sentences to improve readability (B1–B2 level)."
            )

        # 5) Parseability / formatting
        if not breakdown["parseable"] or pf_score < 6:
            fallback.append(
                "Use clean, ATS-friendly formatting (no tables/columns; clear section headings). Export to PDF."
            )

        # 6) Eligibility / location hints
        if elig_loc["score"] < 3:
            fallback.append(
                "Add working rights and current city/country (e.g., “Eligible to work in UK • London”)."
            )

        # De-duplicate and attach
        out["suggestions"] = list(
            dict.fromkeys([s.strip() for s in fallback if s.strip()])
        )

    # 1) Start with LLM/analysis issues + suggestions merged
    out["analysis"]["issues"] = list(dict.fromkeys(
        (out["analysis"]["issues"] or []) + (out.get("suggestions") or [])
    ))
    out["suggestions"] = []  # we no longer ship a separate list
    
    # 2) Add your deterministic fixes (the ones you already craft)
    fixes = []
    if not sec_struct["std"]["experience"]:
        fixes.append("Add a Work Experience section (titles like “Experience” or “Relevant Experience” are fine).")
    elif not sec_struct["reverse_chrono"]:
        fixes.append("Ensure roles are in reverse-chronological order (most recent first) with Month YYYY dates.")
    if breakdown.get("length", 100) < 70:
        fixes.append("Expand to 1–2 pages with impact bullets (8–20 words each).")
    
    # 3) Final de-dupe after appending
    out["analysis"]["issues"] = list(dict.fromkeys((out["analysis"]["issues"] or []) + fixes))

    # ====== NEW: supply a rich issue object + a stringified legacy version ======
    # We’ll derive 'caused_by' from the missing keywords list if available.
    if missing:
        issue = {
            "category": "Keywords",
            "severity": "high",
            "finding": "No domain keywords detected for the target role" if not matched else "Important domain keywords are underrepresented",
            "caused_by": missing[:5],  # e.g., ["Jira", "Scrum"]
            "recommendation": "Add 2–3 relevant tools and frameworks used in your projects",
            "example_fix": "Spearheaded Scrum ceremonies and managed Jira board across 6 squads"
        }
        issue_text = f"{issue['category']} ({issue['severity']}): {issue['finding']} — Fix: {issue['recommendation']} (e.g., {issue['example_fix']})"

        # Keep objects under a new key, and keep strings under the legacy 'issues' key
        a = out.setdefault("analysis", {})
        a.setdefault("issues_objects", [])
        a["issues_objects"].append(issue)

        # Prepend the stringified version so current UI (which renders strings) shows it
        legacy = a.get("issues") or []
        a["issues"] = [issue_text] + legacy

    # ====== END new rich+legacy bridge ======

    # 1) Map various wordings → an intent key and a preferred phrase
    _INTENT_PREF = {
        "formatting_consistency": "Ensure consistent, clean formatting and spacing across all sections.",
        "contact_placement":      "Place contact information at the top of the resume, not inside Experience.",
        "add_certs":              "Add a Certifications section if you hold relevant credentials.",
        "add_experience":         "Add a Work Experience section (e.g., “Experience” or “Relevant Experience”).",
    }
    
    def _intent_key(line: str) -> str | None:
        t = (line or "").lower()
        # formatting
        if re.search(r"\bformat(ting)?\b.*\b(spacing|consistent|consistency)\b", t):
            return "formatting_consistency"
        if re.search(r"\bconsistent\b.*\bformat(ting)?\b", t):
            return "formatting_consistency"
    
        # contact info placement
        if re.search(r"\bcontact\b.*\b(top|header)\b", t):
            return "contact_placement"
        if re.search(r"\bcontact\b.*\bexperience\b", t) and re.search(r"\b(in|inside|embedded|within)\b", t):
            return "contact_placement"
    
        # certifications
        if re.search(r"\b(certification|certifications|credential)\b", t) and re.search(r"\b(add|section)\b", t):
            return "add_certs"
    
        # experience section
        if re.search(r"\b(work )?experience\b.*\b(add|section)\b", t):
            return "add_experience"
        if re.search(r"\badd\b.*\bexperience\b.*\bsection\b", t):
            return "add_experience"
    
        # fall back: None → gets handled by fuzzy step below
        return None
    
    def _canon(s: str) -> str:
        s = (s or "").strip()
        s = re.sub(r"[“”]", '"', s)
        s = re.sub(r"[—–]", "-", s)
        s = re.sub(r"\s+", " ", s)
        return s
    
    def _token_set(s: str) -> set[str]:
        s = _canon(s).lower()
        return set(re.findall(r"[a-z0-9]+", s)) - {"the","a","an","of","and","to","for","with","in","your","resume","section","sections"}
    
    def _similar(a: str, b: str) -> float:
        A, B = _token_set(a), _token_set(b)
        if not A or not B: return 0.0
        inter = len(A & B)
        return inter / float(min(len(A), len(B)))
    
    # 2) Combine issues + (now-empty) suggestions, then reduce by intent (intent wins)
    combined = (out.get("analysis", {}).get("issues") or [])
    
    kept_by_intent: dict[str, str] = {}
    others: list[str] = []
    
    for line in combined:
        if not (line and line.strip()):
            continue
        intent = _intent_key(line)
        if intent:
            kept_by_intent.setdefault(intent, _INTENT_PREF.get(intent) or _canon(line))
        else:
            others.append(_canon(line))
    
    # 3) Fuzzy de-dup for the leftovers (no intent match)
    deduped_others: list[str] = []
    for line in others:
        if not any(_similar(line, x) >= 0.72 or _canon(line) == _canon(x) for x in deduped_others):
            deduped_others.append(line)
    
    # 4) Build the final ordered list: intents in first-seen order, then unique others
    ordered_intents: list[str] = []
    seen_intents = set()
    for line in combined:
        k = _intent_key(line)
        if k and k not in seen_intents and k in kept_by_intent:
            ordered_intents.append(kept_by_intent[k])
            seen_intents.add(k)
    
    final_fixes = ordered_intents + deduped_others
    
    # 5) Write back: one authoritative list, no 'suggestions'
    out.setdefault("analysis", {})["issues"] = final_fixes
    out["suggestions"] = []

    return jsonify(out)

